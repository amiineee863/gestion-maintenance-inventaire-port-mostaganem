from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth import login
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.db.models import Q, Count
from django.core.mail import send_mail
from django.conf import settings
from django.utils import timezone
from datetime import datetime
from django.contrib.auth.views import LoginView
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
from django.core.paginator import Paginator
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from .models import (User, Direction, Bureau, CategorieEquipement, Equipement,
                     DemandeMaintenance, Intervention, PieceRechange, FichierIntervention, LogAction)
from .forms import (UserRegistrationForm, EquipementForm, DemandeMaintenanceForm,
                    AssignationTechnicienForm, InterventionForm, PieceRechangeFormSet,
                    FiltreDemandeForm, FiltreEquipementForm, FichierInterventionFormSet, FiltreLogForm, FiltreInterventionForm)


# ============= HELPERS =============
def log_action(user, action, type_objet='', objet_id=None, details='', request=None):
    """Fonction utilitaire pour créer un log d'action"""
    ip_address = None
    if request:
        x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            ip_address = x_forwarded_for.split(',')[0]
        else:
            ip_address = request.META.get('REMOTE_ADDR')
    
    LogAction.objects.create(
        utilisateur=user,
        action=action,
        type_objet=type_objet,
        objet_id=objet_id,
        details=details,
        adresse_ip=ip_address
    )
def is_admin(user):
    return user.is_authenticated and user.role == 'ADMIN'

def is_technicien(user):
    return user.is_authenticated and user.role == 'TECHNICIEN'

def is_employe(user):
    return user.is_authenticated and user.role == 'EMPLOYE'


def envoyer_email_notification(demande):
    """Envoie un email de notification à l'employé"""
    if not demande.email_envoye and demande.employe.email:
        try:
            subject = f"Demande #{demande.pk} - Réparation terminée"
            message = f"""
Bonjour {demande.employe.get_full_name()},

Votre demande de maintenance #{demande.pk} pour l'équipement {demande.equipement.code_equipement} 
a été traitée et est maintenant terminée.

Merci de vous connecter pour valider ou signaler un problème.

Cordialement,
Service Maintenance EP Mostaganem
            """
            send_mail(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,
                [demande.employe.email],
                fail_silently=False,
            )
            demande.email_envoye = True
            demande.save()
        except Exception as e:
            print(f"Erreur envoi email: {e}")


# ============= VUES COMMUNES =============

def home(request):
    """Page d'accueil - redirige selon le rôle"""
    if not request.user.is_authenticated:
        return render(request, 'maintenance/home.html')
    
    if request.user.role == 'ADMIN':
        return redirect('admin_dashboard')
    elif request.user.role == 'TECHNICIEN':
        return redirect('technicien_dashboard')
    else:
        return redirect('employe_dashboard')
    

class CustomLoginView(LoginView):
    """Vue de login personnalisée qui redirige si déjà connecté"""
    template_name = 'maintenance/login.html'
    
    def dispatch(self, request, *args, **kwargs):
        # Si l'utilisateur est déjà connecté, le rediriger vers son dashboard
        if request.user.is_authenticated:
            if request.user.role == 'ADMIN':
                return redirect('admin_dashboard')
            elif request.user.role == 'TECHNICIEN':
                return redirect('technicien_dashboard')
            else:
                return redirect('employe_dashboard')
        return super().dispatch(request, *args, **kwargs)


# ============= ESPACE EMPLOYÉ =============

@login_required
@user_passes_test(is_employe)
def employe_dashboard(request):
    """Tableau de bord de l'employé"""
    demandes = DemandeMaintenance.objects.filter(employe=request.user).select_related(
        'equipement', 'technicien'
    ).order_by('-date_creation')
    
    context = {
        'demandes': demandes,
        'total': demandes.count(),
        'en_attente': demandes.filter(statut='EN_ATTENTE').count(),
        'en_cours': demandes.filter(statut__in=['ASSIGNEE', 'EN_COURS']).count(),
        'terminees': demandes.filter(statut='TERMINEE').count(),
    }
    return render(request, 'maintenance/employe/dashboard.html', context)


@login_required
@user_passes_test(is_employe)
def employe_creer_demande(request):
    """Création d'une nouvelle demande de maintenance"""
    if request.method == 'POST':
        form = DemandeMaintenanceForm(request.POST, user=request.user)  # ← AJOUT user=request.user
        if form.is_valid():
            demande = form.save(commit=False)
            demande.employe = request.user
            demande.save()
            log_action(
                user=request.user,
                action='DEMANDE_CREATION',
                type_objet='DemandeMaintenance',
                objet_id=demande.pk,
                details=f"Création demande pour équipement {demande.equipement.code_equipement}",
                request=request
            )
            messages.success(request, 'Demande de maintenance créée avec succès.')
            return redirect('employe_dashboard')
    else:
        form = DemandeMaintenanceForm(user=request.user)  # ← AJOUT user=request.user
    
    return render(request, 'maintenance/employe/creer_demande.html', {'form': form})


@login_required
@user_passes_test(is_employe)
def employe_modifier_demande(request, pk):
    """Modification d'une demande (seulement si EN_ATTENTE)"""
    demande = get_object_or_404(DemandeMaintenance, pk=pk, employe=request.user)
    
    if not demande.peut_etre_modifiee():
        messages.error(request, 'Cette demande ne peut plus être modifiée.')
        return redirect('employe_dashboard')
    
    if request.method == 'POST':
        form = DemandeMaintenanceForm(request.POST, instance=demande, user=request.user)  # ← AJOUT user=request.user
        if form.is_valid():
            form.save()
            log_action(
                user=request.user,
                action='DEMANDE_MODIFICATION',
                type_objet='DemandeMaintenance',
                objet_id=demande.pk,
                details=f"Modification demande #{demande.pk}",
                request=request
            )
            messages.success(request, 'Demande modifiée avec succès.')
            return redirect('employe_dashboard')
    else:
        # Pré-remplir avec le code de l'équipement actuel
        initial_data = {'code_equipement': demande.equipement.code_equipement}  # ← AJOUT DE CETTE LIGNE
        form = DemandeMaintenanceForm(instance=demande, initial=initial_data, user=request.user)  # ← AJOUT initial=initial_data, user=request.user
    
    return render(request, 'maintenance/employe/modifier_demande.html', {'form': form, 'demande': demande})


@login_required
@user_passes_test(is_employe)
def employe_supprimer_demande(request, pk):
    """Suppression d'une demande (seulement si EN_ATTENTE)"""
    demande = get_object_or_404(DemandeMaintenance, pk=pk, employe=request.user)
    demande_id = demande.pk
    equipement_id = demande.equipement.pk
    
    if not demande.peut_etre_supprimee():
        messages.error(request, 'Cette demande ne peut plus être supprimée.')
        return redirect('employe_dashboard')
    
    if request.method == 'POST':
        log_action(
            user=request.user,
            action='DEMANDE_SUPPRESSION',
            type_objet='DemandeMaintenance',
            objet_id=demande_id,
            details=f"Suppression demande #{demande_id} - {equipement_id}",
            request=request
        )
        demande.delete()
        messages.success(request, 'Demande supprimée avec succès.')
        return redirect('employe_dashboard')
        
        
    
    return render(request, 'maintenance/employe/supprimer_demande.html', {'demande': demande})


@login_required
@user_passes_test(is_employe)
def employe_valider_demande(request, pk):
    """Validation de la réparation par l'employé"""
    demande = get_object_or_404(DemandeMaintenance, pk=pk, employe=request.user)
    
    if not demande.peut_etre_validee():
        messages.error(request, 'Cette demande ne peut pas être validée.')
        return redirect('employe_dashboard')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'valider':
            demande.statut = 'VALIDEE'
            log_action(
                user=request.user,
                action='DEMANDE_VALIDATION',
                type_objet='DemandeMaintenance',
                objet_id=demande.pk,
                details=f"Validation réparation demande #{demande.pk}",
                request=request
            )
            messages.success(request, 'Demande validée. Merci pour votre retour.')
        elif action == 'refuser':
            demande.statut = 'REFUSEE'
            log_action(
                user=request.user,
                action='DEMANDE_REFUS',
                type_objet='DemandeMaintenance',
                objet_id=demande.pk,
                details=f"Refus réparation demande #{demande.pk}",
                request=request
            )
            messages.warning(request, 'Problème signalé. Un technicien sera informé.')
        demande.save()
        return redirect('employe_dashboard')
    
    return render(request, 'maintenance/employe/valider_demande.html', {'demande': demande})


# ============= ESPACE TECHNICIEN =============

@login_required
@user_passes_test(is_technicien)
def technicien_dashboard(request):
    """Tableau de bord du technicien"""
    demandes = DemandeMaintenance.objects.filter(
        technicien=request.user
    ).select_related('equipement', 'employe').order_by('-date_creation')
    
    context = {
        'demandes': demandes,
        'total': demandes.count(),
        'assignees': demandes.filter(statut='ASSIGNEE').count(),
        'en_cours': demandes.filter(statut='EN_COURS').count(),
        'terminees': demandes.filter(statut='TERMINEE').count(),
    }
    return render(request, 'maintenance/technicien/dashboard.html', context)


@login_required
@user_passes_test(is_technicien)
def technicien_detail_demande(request, pk):
    """Détail et traitement d'une demande"""
    demande = get_object_or_404(DemandeMaintenance, pk=pk, technicien=request.user)
    
    # Récupérer l'intervention si elle existe
    try:
        intervention = demande.intervention
    except Intervention.DoesNotExist:
        intervention = None
    
    context = {
        'demande': demande,
        'intervention': intervention,
    }
    return render(request, 'maintenance/technicien/detail_demande.html', context)


@login_required
@user_passes_test(is_technicien)
def technicien_changer_statut(request, pk):
    """Changer le statut d'une demande"""
    demande = get_object_or_404(DemandeMaintenance, pk=pk, technicien=request.user)
    
    if request.method == 'POST':
        nouveau_statut = request.POST.get('statut')
        
        if nouveau_statut in ['ASSIGNEE', 'EN_COURS', 'TERMINEE']:
            demande.statut = nouveau_statut
            demande.save()
            log_action(
                user=request.user,
                action='STATUT_CHANGE',
                type_objet='DemandeMaintenance',
                objet_id=demande.pk,
                details=f"Changement statut demande #{demande.pk} → {demande.get_statut_display()}",
                request=request
            )
            
            # Envoyer email si terminée
            if nouveau_statut == 'TERMINEE':
                envoyer_email_notification(demande)
            
            messages.success(request, f'Statut changé à : {demande.get_statut_display()}')
        
        return redirect('technicien_detail_demande', pk=pk)
    
    return redirect('technicien_dashboard')

@login_required
@user_passes_test(is_technicien)
def technicien_creer_intervention(request, pk):
    demande = get_object_or_404(
        DemandeMaintenance,
        pk=pk,
        technicien=request.user
    )

    if hasattr(demande, 'intervention'):
        messages.warning(request, "Une intervention existe déjà.")
        return redirect('technicien_modifier_intervention', pk=pk)

    if request.method == 'POST':
        form = InterventionForm(request.POST)

        piece_formset = PieceRechangeFormSet(
            request.POST,
            prefix='pieces'
        )

        fichier_formset = FichierInterventionFormSet(
            request.POST,
            request.FILES,
            prefix='fichiers'
        )

        if (
            form.is_valid()
            and piece_formset.is_valid()
            and fichier_formset.is_valid()
        ):
            intervention = form.save(commit=False)
            intervention.demande = demande
            intervention.save()

            piece_formset.instance = intervention
            piece_formset.save()

            fichier_formset.instance = intervention
            fichier_formset.save()

            log_action(
                user=request.user,
                action='INTERVENTION_CREATION',
                type_objet='Intervention',
                objet_id=intervention.pk,
                details=f"Création rapport pour demande #{demande.pk}",
                request=request
            )

            messages.success(request, "Rapport d'intervention créé avec succès.")
            return redirect('technicien_detail_demande', pk=pk)

    else:
        form = InterventionForm()
        piece_formset = PieceRechangeFormSet(prefix='pieces')
        fichier_formset = FichierInterventionFormSet(prefix='fichiers')

    return render(request, 'maintenance/technicien/creer_intervention.html', {
        'form': form,
        'formset': piece_formset,
        'fichier_formset': fichier_formset,
        'demande': demande,
    })



@login_required
@user_passes_test(is_technicien)
def technicien_modifier_intervention(request, pk):
    """Modification d'un rapport d'intervention"""
    demande = get_object_or_404(DemandeMaintenance, pk=pk, technicien=request.user)
    intervention = get_object_or_404(Intervention, demande=demande)
    
    if request.method == 'POST':
        form = InterventionForm(request.POST, instance=intervention)
        formset = PieceRechangeFormSet(request.POST, instance=intervention)
        
        if form.is_valid() and formset.is_valid():
            form.save()
            formset.save()

            log_action(
                user=request.user,
                action='INTERVENTION_MODIFICATION',
                type_objet='Intervention',
                objet_id=intervention.pk,
                details=f"Modification rapport intervention #{intervention.pk}",
                request=request
            )
            
            messages.success(request, 'Rapport d\'intervention modifié avec succès.')
            return redirect('technicien_detail_demande', pk=pk)
    else:
        form = InterventionForm(instance=intervention)
        formset = PieceRechangeFormSet(instance=intervention)
    
    return render(request, 'maintenance/technicien/modifier_intervention.html', {
        'form': form,
        'formset': formset,
        'demande': demande,
        'intervention': intervention,
    })


# ============= ESPACE ADMIN =============

@login_required
@user_passes_test(is_admin)
def admin_dashboard(request):
    """Tableau de bord administrateur avec statistiques"""
    # Statistiques générales
    total_demandes = DemandeMaintenance.objects.count()
    total_equipements = Equipement.objects.count()
    total_techniciens = User.objects.filter(role='TECHNICIEN', is_active=True).count()
    
    # Demandes par statut
    demandes_par_statut = DemandeMaintenance.objects.values('statut').annotate(count=Count('id'))
    
    # Demandes récentes
    demandes_recentes = DemandeMaintenance.objects.select_related(
        'equipement', 'employe', 'technicien'
    ).order_by('-date_creation')[:10]
    
    # Pannes par marque (top 5)
    pannes_par_marque = Equipement.objects.values('marque').annotate(
        nb_pannes=Count('demandes')
    ).order_by('-nb_pannes')[:5]
    
    context = {
        'total_demandes': total_demandes,
        'total_equipements': total_equipements,
        'total_techniciens': total_techniciens,
        'demandes_par_statut': demandes_par_statut,
        'demandes_recentes': demandes_recentes,
        'pannes_par_marque': pannes_par_marque,
    }
    return render(request, 'maintenance/admin/dashboard.html', context)


@login_required
@user_passes_test(is_admin)
def admin_liste_demandes(request):
    """Liste des demandes avec filtres"""
    demandes = DemandeMaintenance.objects.select_related(
        'equipement', 'employe', 'technicien'
    ).order_by('-date_creation')
    
    # Appliquer les filtres
    form = FiltreDemandeForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('statut'):
            demandes = demandes.filter(statut=form.cleaned_data['statut'])
        if form.cleaned_data.get('urgence'):
            demandes = demandes.filter(urgence=form.cleaned_data['urgence'])
        if form.cleaned_data.get('technicien'):
            demandes = demandes.filter(technicien=form.cleaned_data['technicien'])
        if form.cleaned_data.get('date_debut'):
            demandes = demandes.filter(date_creation__gte=form.cleaned_data['date_debut'])
        if form.cleaned_data.get('date_fin'):
            demandes = demandes.filter(date_creation__lte=form.cleaned_data['date_fin'])
        if form.cleaned_data.get('code_equipement'):
            demandes = demandes.filter(equipement__code_equipement__icontains=form.cleaned_data['code_equipement'])
        if form.cleaned_data.get('categorie'):
            demandes = demandes.filter(equipement__categorie=form.cleaned_data['categorie'])

    
    total_affiche = demandes.count()

    en_attente = demandes.filter(statut='EN_ATTENTE').count()
    en_cours = demandes.filter(statut='EN_COURS').count()
    terminees = demandes.filter(statut='TERMINEE').count()
    validees = demandes.filter(statut='VALIDEE').count()
    refusees = demandes.filter(statut='REFUSEE').count()

    # Pagination
    paginator = Paginator(demandes, 10)  # 10 demandes par page
    page_number = request.GET.get('page')
    demandes_page = paginator.get_page(page_number)

    context = {
        'demandes': demandes_page,
        'form': form,
        'total_affiche': total_affiche,
        'en_attente': en_attente,
        'en_cours': en_cours,
        'terminees': terminees,
        'validees': validees,
        'refusees' : refusees,
    }

    
    return render(request, 'maintenance/admin/liste_demandes.html', context)


@login_required
@user_passes_test(is_admin)
def admin_assigner_technicien(request, pk):
    """Assigner un technicien à une demande"""
    demande = get_object_or_404(DemandeMaintenance, pk=pk)
    
    if request.method == 'POST':
        form = AssignationTechnicienForm(request.POST, instance=demande)
        if form.is_valid():
            demande = form.save(commit=False)
            demande.statut = 'ASSIGNEE'
            demande.save()
            log_action(
                user=request.user,
                action='ASSIGNATION',
                type_objet='DemandeMaintenance',
                objet_id=demande.pk,
                details=f"Assignation {demande.technicien.get_full_name()} à demande #{demande.pk}",
                request=request
            )
            messages.success(request, f'Demande assignée à {demande.technicien.get_full_name()}')
            return redirect('admin_liste_demandes')
    else:
        form = AssignationTechnicienForm(instance=demande)
    
    return render(request, 'maintenance/admin/assigner_technicien.html', {
        'form': form,
        'demande': demande,
    })


@login_required
@user_passes_test(is_admin)
def admin_liste_equipements(request):
    """Liste des équipements avec filtres"""
    equipements = Equipement.objects.select_related('bureau', 'categorie').order_by('code_equipement')
    
    # Appliquer les filtres
    form = FiltreEquipementForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('code_equipement'):
            equipements = equipements.filter(code_equipement__icontains=form.cleaned_data['code_equipement'])
        if form.cleaned_data.get('marque'):
            equipements = equipements.filter(marque__icontains=form.cleaned_data['marque'])
        if form.cleaned_data.get('categorie'):
            equipements = equipements.filter(categorie=form.cleaned_data['categorie'])
        if form.cleaned_data.get('bureau'):
            equipements = equipements.filter(bureau=form.cleaned_data['bureau'])
    
    # Pagination
    paginator = Paginator(equipements, 10)  # 10 équipements par page
    page_number = request.GET.get('page')
    equipements_page = paginator.get_page(page_number)
    context = {
        'equipements': equipements_page,
        'form': form,
    }
    return render(request, 'maintenance/admin/liste_equipements.html', context)


@login_required
@user_passes_test(is_admin)
def admin_creer_equipement(request):
    """Création d'un nouvel équipement"""
    if request.method == 'POST':
        form = EquipementForm(request.POST)
        if form.is_valid():
            equipement = form.save()
            log_action(
                user=request.user,
                action='EQUIPEMENT_CREATION',
                type_objet='Equipement',
                objet_id=None,
                details=f"Création équipement {equipement.code_equipement}",
                request=request
            )
            messages.success(request, 'Équipement créé avec succès.')
            
            return redirect('admin_liste_equipements')
    else:
        form = EquipementForm()
    
    return render(request, 'maintenance/admin/creer_equipement.html', {'form': form})


@login_required
@user_passes_test(is_admin)
def admin_modifier_equipement(request, code):
    """Modification d'un équipement"""
    equipement = get_object_or_404(Equipement, code_equipement=code)
    
    if request.method == 'POST':
        form = EquipementForm(request.POST, instance=equipement)
        if form.is_valid():
            form.save()

            log_action(
                user=request.user,
                action='EQUIPEMENT_MODIFICATION',
                type_objet='Equipement',
                objet_id=None,
                details=f"Modification équipement {equipement.code_equipement}",
                request=request
            )
            messages.success(request, 'Équipement modifié avec succès.')
            return redirect('admin_liste_equipements')
    else:
        form = EquipementForm(instance=equipement)
    
    return render(request, 'maintenance/admin/modifier_equipement.html', {
        'form': form,
        'equipement': equipement,
    })


@login_required
@user_passes_test(is_admin)
def admin_supprimer_equipement(request, code):
    """Suppression d'un équipement"""
    equipement = get_object_or_404(Equipement, code_equipement=code)
    if request.method == 'POST':
        equipement.delete()
        log_action(
            user=request.user,
            action='EQUIPEMENT_SUPPRESSION',
            type_objet='Equipement',
            objet_id=None,
            details=f"Suppression équipement {code}",
            request=request
        )
        messages.success(request, 'Équipement supprimé avec succès.')
        return redirect('admin_liste_equipements')
    
    return render(request, 'maintenance/admin/supprimer_equipement.html', {'equipement': equipement})

@login_required
@user_passes_test(is_admin)
def admin_import_equipements_csv(request):
    """Importer des équipements depuis un fichier CSV"""
    if request.method == 'POST' and request.FILES.get('csv_file'):
        csv_file = request.FILES['csv_file']
        
        if not csv_file.name.endswith('.csv'):
            messages.error(request, 'Le fichier doit être au format CSV.')
            return redirect('admin_import_equipements_csv')
        
        try:
            import csv
            from io import StringIO
            
            # Lire le fichier CSV
            decoded_file = csv_file.read().decode('utf-8')
            io_string = StringIO(decoded_file)
            reader = csv.DictReader(io_string)
            
            count_created = 0
            count_updated = 0
            errors = []
            
            for row in reader:
                try:
                    # Récupérer ou créer le bureau
                    bureau = None
                    if row.get('bureau') and row.get('direction'):
                        direction, _ = Direction.objects.get_or_create(nom=row['direction'])
                        bureau, _ = Bureau.objects.get_or_create(nom=row['bureau'], direction=direction)
                    
                    # Récupérer ou créer la catégorie
                    categorie = None
                    if row.get('categorie'):
                        categorie, _ = CategorieEquipement.objects.get_or_create(nom=row['categorie'])
                    
                    # Créer ou mettre à jour l'équipement
                    equipement, created = Equipement.objects.update_or_create(
                        code_equipement=row['code_equipement'],
                        defaults={
                            'nom': row['nom'],
                            'marque': row['marque'],
                            'date_acquisition': row['date_acquisition'],
                            'description_technique': row.get('description_technique', ''),
                            'bureau': bureau,
                            'categorie': categorie,
                        }
                    )
                    
                    if created:
                        count_created += 1
                    else:
                        count_updated += 1
                        
                except Exception as e:
                    errors.append(f"Ligne {reader.line_num}: {str(e)}")
            
            if errors:
                for error in errors[:5]:  # Montrer max 5 erreurs
                    messages.warning(request, error)
            
            log_action(
                user=request.user,
                action='IMPORT_CSV',
                details=f"Import CSV: {count_created} créés, {count_updated} mis à jour",
                request=request
            )
            messages.success(request, f'{count_created} équipements créés, {count_updated} mis à jour.')
            return redirect('admin_liste_equipements')
            
        except Exception as e:
            messages.error(request, f'Erreur lors de l\'import: {str(e)}')
    
    return render(request, 'maintenance/admin/import_equipements.html')
# ============= EXPORTS =============

@login_required
@user_passes_test(is_admin)
def export_demandes_csv(request):
    """Export des demandes en CSV"""
    demandes = DemandeMaintenance.objects.select_related('equipement', 'employe', 'technicien').order_by('-date_creation')
    
    # Appliquer les mêmes filtres que la liste
    form = FiltreDemandeForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('statut'):
            demandes = demandes.filter(statut=form.cleaned_data['statut'])
        if form.cleaned_data.get('urgence'):
            demandes = demandes.filter(urgence=form.cleaned_data['urgence'])
        if form.cleaned_data.get('technicien'):
            demandes = demandes.filter(technicien=form.cleaned_data['technicien'])
        if form.cleaned_data.get('date_demande'):
            demandes = demandes.filter(date_creation__date=form.cleaned_data['date_demande'])
        if form.cleaned_data.get('code_equipement'):
            demandes = demandes.filter(equipement__code_equipement__icontains=form.cleaned_data['code_equipement'])
        if form.cleaned_data.get('categorie'):
            demandes = demandes.filter(equipement__categorie=form.cleaned_data['categorie'])
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="demandes_{datetime.now().strftime("%Y%m%d")}.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['ID', 'Date', 'Équipement', 'Employé', 'Technicien', 'Urgence', 'Statut', 'Description'])
    
    
    for d in demandes:
        writer.writerow([
            d.pk,
            d.date_creation.strftime('%Y-%m-%d %H:%M'),
            d.equipement.code_equipement,
            d.employe.get_full_name(),
            d.technicien.get_full_name() if d.technicien else 'Non assigné',
            d.get_urgence_display(),
            d.get_statut_display(),
            d.description[:100],
        ])
    
    log_action(
        user=request.user,
        action='EXPORT_CSV',
        details="Export CSV liste demandes",
        request=request
    )
    
    return response


@login_required
@user_passes_test(is_admin)
def export_demandes_pdf(request):
    """Export des demandes en PDF"""
    demandes = DemandeMaintenance.objects.select_related('equipement', 'employe', 'technicien').order_by('-date_creation')
    
    # Appliquer les mêmes filtres que la liste
    form = FiltreDemandeForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('statut'):
            demandes = demandes.filter(statut=form.cleaned_data['statut'])
        if form.cleaned_data.get('urgence'):
            demandes = demandes.filter(urgence=form.cleaned_data['urgence'])
        if form.cleaned_data.get('technicien'):
            demandes = demandes.filter(technicien=form.cleaned_data['technicien'])
        if form.cleaned_data.get('date_demande'):
            demandes = demandes.filter(date_creation__date=form.cleaned_data['date_demande'])
        if form.cleaned_data.get('code_equipement'):
            demandes = demandes.filter(equipement__code_equipement__icontains=form.cleaned_data['code_equipement'])
        if form.cleaned_data.get('categorie'):
            demandes = demandes.filter(equipement__categorie=form.cleaned_data['categorie'])
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
    elements = []
    
    styles = getSampleStyleSheet()
    title = Paragraph("Rapport des Demandes de Maintenance", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 20))
    
    # Données
    data = [['ID', 'Date', 'Équipement', 'Employé', 'Technicien', 'Urgence', 'Statut']]
   
    
    for d in demandes:
        data.append([
            str(d.pk),
            d.date_creation.strftime('%Y-%m-%d'),
            d.equipement.code_equipement,
            d.employe.get_full_name(),
            d.technicien.get_full_name() if d.technicien else 'N/A',
            d.get_urgence_display(),
            d.get_statut_display(),
        ])
    
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    
    elements.append(table)
    doc.build(elements)
    
    buffer.seek(0)

    log_action(
        user=request.user,
        action='EXPORT_PDF',
        details="Export PDF liste demandes",
        request=request
    )
    return HttpResponse(buffer, content_type='application/pdf')


@login_required
@user_passes_test(is_admin)
def export_equipements_csv(request):
    """Export des équipements en CSV"""
    equipements = Equipement.objects.select_related('bureau__direction', 'categorie').order_by('code_equipement')
    
    # Appliquer les mêmes filtres que la liste
    form = FiltreEquipementForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('marque'):
            equipements = equipements.filter(marque__icontains=form.cleaned_data['marque'])
        if form.cleaned_data.get('categorie'):
            equipements = equipements.filter(categorie=form.cleaned_data['categorie'])
        if form.cleaned_data.get('bureau'):
            equipements = equipements.filter(bureau=form.cleaned_data['bureau'])
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="equipements_{datetime.now().strftime("%Y%m%d")}.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Code', 'Nom', 'Marque', 'Catégorie', 'Bureau', 'Direction', 'Date Acquisition'])
    for e in equipements:
        writer.writerow([
            e.code_equipement,
            e.nom,
            e.marque,
            e.categorie.nom if e.categorie else '',
            e.bureau.nom if e.bureau else '',
            e.bureau.direction.nom if e.bureau else '',
            e.date_acquisition.strftime('%Y-%m-%d'),
        ])
    
    log_action(
        user=request.user,
        action='EXPORT_CSV',
        details="Export CSV liste equipements",
        request=request
    )
    
    return response


# ============= CONSULTATION INTERVENTIONS (ADMIN) =============

@login_required
@user_passes_test(is_admin)
def admin_liste_interventions(request):
    """Liste de toutes les interventions avec leurs fichiers"""
    interventions = Intervention.objects.select_related(
        'demande__equipement', 'demande__employe', 'demande__technicien'
    ).prefetch_related('fichiers', 'pieces').order_by('-date_intervention')
    
    # Appliquer les filtres
    form = FiltreInterventionForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('technicien'):
            interventions = interventions.filter(demande__technicien=form.cleaned_data['technicien'])
        if form.cleaned_data.get('code_equipement'):
            interventions = interventions.filter(demande__equipement__code_equipement__icontains=form.cleaned_data['code_equipement'])
        if form.cleaned_data.get('type_reparation'):
            interventions = interventions.filter(type_reparation=form.cleaned_data['type_reparation'])
        if form.cleaned_data.get('date_debut'):
            interventions = interventions.filter(date_intervention__gte=form.cleaned_data['date_debut'])
        if form.cleaned_data.get('date_fin'):
            from datetime import datetime, time
            date_fin = datetime.combine(form.cleaned_data['date_fin'], time.max)
            interventions = interventions.filter(date_intervention__lte=date_fin)
    
    total_interventions = interventions.count()

    reparations_internes = interventions.filter(
        type_reparation='INTERNE'
    ).count()

    reparations_externes = interventions.filter(
        type_reparation='EXTERNE'
    ).count()

    documents_joints = interventions.filter(
        fichiers__isnull=False
    ).distinct().count()

    # Pagination
    paginator = Paginator(interventions, 10)  
    page_number = request.GET.get('page')
    interventions_page = paginator.get_page(page_number)

    context = {
        'interventions': interventions_page,
        'form': form,
        'total_interventions': total_interventions,
        'reparations_internes': reparations_internes,
        'reparations_externes': reparations_externes,
        'documents_joints': documents_joints,
    }
    return render(request, 'maintenance/admin/liste_interventions.html', context)

@login_required
@user_passes_test(is_admin)
def admin_detail_intervention(request, pk):
    """Détail d'une intervention avec tous ses fichiers"""
    intervention = get_object_or_404(
        Intervention.objects.select_related(
            'demande__equipement__bureau__direction',
            'demande__employe',
            'demande__technicien'
        ).prefetch_related('fichiers', 'pieces'),
        pk=pk
    )
    
    # Grouper les fichiers par type
    fichiers_par_type = {}
    for fichier in intervention.fichiers.all():
        type_fichier = fichier.get_type_fichier_display()
        if type_fichier not in fichiers_par_type:
            fichiers_par_type[type_fichier] = []
        fichiers_par_type[type_fichier].append(fichier)
    
    context = {
        'intervention': intervention,
        'demande': intervention.demande,
        'fichiers_par_type': fichiers_par_type,
    }
    return render(request, 'maintenance/admin/detail_intervention.html', context)


@login_required
@user_passes_test(is_admin)
def admin_supprimer_fichier(request, pk):
    """Supprimer un fichier d'intervention"""
    fichier = get_object_or_404(FichierIntervention, pk=pk)
    intervention_pk = fichier.intervention.pk
    
    if request.method == 'POST':
        fichier.fichier.delete()  # Supprimer le fichier physique
        fichier.delete()  # Supprimer l'entrée en base

        log_action(
            user=request.user,
            action='FICHIER_SUPPRESSION',
            type_objet='Intervention',
            objet_id=intervention_pk,
            details=f"Suppression fichier intervention #{intervention_pk}",
            request=request
        )
        messages.success(request, 'Fichier supprimé avec succès.')
        return redirect('admin_detail_intervention', pk=intervention_pk)
    
    return render(request, 'maintenance/admin/supprimer_fichier.html', {'fichier': fichier})



@login_required
def export_intervention_pdf(request, pk):
    """Export d'une intervention en PDF"""
    # Vérifier les permissions
    intervention = get_object_or_404(Intervention, pk=pk)
    user = request.user
    
    # Admin ou technicien de l'intervention
    if not (user.role == 'ADMIN' or (user.role == 'TECHNICIEN' and intervention.demande.technicien == user)):
        messages.error(request, 'Vous n\'avez pas la permission d\'exporter ce rapport.')
        return redirect('home')
    
    log_action(
        user=request.user,
        action='EXPORT_PDF',
        type_objet='Intervention',
        objet_id=intervention.pk,
        details=f"Export PDF intervention #{intervention.pk}",
        request=request
    )
    
    
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    elements = []
    styles = getSampleStyleSheet()
    
    # Titre
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#0d6efd'),
        spaceAfter=20,
        spaceBefore=20,
        alignment=1  # Center
    )
    header_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading3'],
        spaceAfter=20,
        alignment=1  # Center
    )
    
    elements.append(Paragraph("REPUBLIQUE ALGERIENNE DEMOCRATIQUE POPULAIRE", header_style))
    elements.append(Paragraph("MINISTERE DES TRANSPORTS", styles['Heading4']))
    elements.append(Paragraph("GROUPE SERVICES PORTUAIRES « SERPORT SPA»", styles['Heading4']))
    elements.append(Paragraph("ENTREPRISE PORTUAIRE DE MOSTAGANEM", styles['Heading4']))
    elements.append(Paragraph("RAPPORT TECHNIQUE", title_style))
    elements.append(Paragraph(f"Intervention #{intervention.pk}", styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    # Informations générales
    data = [
        ['Date d\'intervention:', intervention.date_intervention.strftime('%d/%m/%Y à %H:%M')],
        ['Type de réparation:', intervention.get_type_reparation_display()],
        ['Demande associée:', f"#{intervention.demande.pk}"],
        ['Demandeur:', intervention.demande.employe.get_full_name()],
        ['Technicien:', intervention.demande.technicien.get_full_name()],
    ]
    
    table = Table(data, colWidths=[5*cm, 12*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 10))
    
    # Équipement
    elements.append(Paragraph("ÉQUIPEMENT CONCERNÉ", styles['Heading3']))
    data_eq = [
        ['Code:', intervention.demande.equipement.code_equipement],
        ['Nom:', intervention.demande.equipement.nom],
        ['Marque:', intervention.demande.equipement.marque],
        ['Catégorie:', intervention.demande.equipement.categorie.nom if intervention.demande.equipement.categorie else '-'],
        ['Bureau:', intervention.demande.equipement.bureau.nom if intervention.demande.equipement.bureau else '-'],
    ]
    table_eq = Table(data_eq, colWidths=[5*cm, 12*cm])
    table_eq.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
    ]))
    elements.append(table_eq)
    elements.append(Spacer(1, 10))
    
    # Détails intervention
    elements.append(Paragraph("DÉTAILS DE L'INTERVENTION", styles['Heading3']))
    details_style = ParagraphStyle('Details', parent=styles['BodyText'], fontSize=10, leading=14)
    elements.append(Paragraph(intervention.details.replace('\n', '<br/>'), details_style))
    elements.append(Spacer(1, 10))
    
    # Pièces
    if intervention.pieces.exists():
        elements.append(Paragraph("PIÈCES DE RECHANGE UTILISÉES", styles['Heading3']))
        pieces_data = [['Nom', 'Prix Unitaire', 'Quantité', 'Total']]
        for piece in intervention.pieces.all():
            pieces_data.append([
                piece.nom,
                f"{piece.prix_unitaire} DA",
                str(piece.quantite),
                f"{piece.cout_total()} DA"
            ])
        pieces_data.append(['', '', 'TOTAL:', f"{intervention.cout_total_pieces()} DA"])
        
        table_pieces = Table(pieces_data, colWidths=[8*cm, 3*cm, 3*cm, 3*cm])
        table_pieces.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, -1), (-1, -1), colors.lightgreen),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ]))
        elements.append(table_pieces)
        elements.append(Spacer(1, 10))


    data = [["VISA DE L'INFORMATICIEN", "VISA DE RESPONSABLE DE LA CELLULE"]]
    # On définit la largeur des colonnes (ex: 250 points chacune)
    t = Table(data, colWidths=[250, 250])

    t.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
    ]))

    elements.append(t)
    
    
    doc.build(elements)
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="intervention_{intervention.pk}.pdf"'
    return response


@login_required
def export_intervention_word(request, pk):
    """Export d'une intervention en Word"""
    intervention = get_object_or_404(Intervention, pk=pk)
    user = request.user
    
    # Vérifier permissions
    if not (user.role == 'ADMIN' or (user.role == 'TECHNICIEN' and intervention.demande.technicien == user)):
        messages.error(request, 'Vous n\'avez pas la permission d\'exporter ce rapport.')
        return redirect('home')
    
    log_action(
        user=request.user,
        action='EXPORT_WORD',
        type_objet='Intervention',
        objet_id=intervention.pk,
        details=f"Export Word intervention #{intervention.pk}",
        request=request
    )
    
    # Créer document Word
    document = Document()

    # --- EN-TÊTE ---
    # République (Centré)
    header_rep = document.add_paragraph("REPUBLIQUE ALGERIENNE DEMOCRATIQUE POPULAIRE")
    header_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_rep.runs[0].bold = True

    # Ministères et Entreprise (Aligné à gauche)
    header_min = document.add_paragraph("MINISTERE DES TRANSPORTS\n"
                                        "GROUPE SERVICES PORTUAIRES « SERPORT SPA»\n"
                                        "ENTREPRISE PORTUAIRE DE MOSTAGANEM")
    header_min.alignment = WD_ALIGN_PARAGRAPH.LEFT

    
    # Titre
    title = document.add_heading(f'RAPPORT TECHNIQUE D\'INTERVENTION #{intervention.pk}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informations
    document.add_heading('Informations Générales', level=1)
    table = document.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    table.rows[0].cells[0].text = 'Date d\'intervention:'
    table.rows[0].cells[1].text = intervention.date_intervention.strftime('%d/%m/%Y à %H:%M')
    table.rows[1].cells[0].text = 'Type de réparation:'
    table.rows[1].cells[1].text = intervention.get_type_reparation_display()
    table.rows[2].cells[0].text = 'Demande:'
    table.rows[2].cells[1].text = f"#{intervention.demande.pk}"
    table.rows[3].cells[0].text = 'Demandeur:'
    table.rows[3].cells[1].text = intervention.demande.employe.get_full_name()
    table.rows[4].cells[0].text = 'Technicien:'
    table.rows[4].cells[1].text = intervention.demande.technicien.get_full_name()
    
    # Équipement
    document.add_heading('Équipement Concerné', level=1)
    eq_table = document.add_table(rows=5, cols=2)
    eq_table.style = 'Light Grid Accent 1'
    eq_table.rows[0].cells[0].text = 'Code:'
    eq_table.rows[0].cells[1].text = intervention.demande.equipement.code_equipement
    eq_table.rows[1].cells[0].text = 'Nom:'
    eq_table.rows[1].cells[1].text = intervention.demande.equipement.nom
    eq_table.rows[2].cells[0].text = 'Marque:'
    eq_table.rows[2].cells[1].text = intervention.demande.equipement.marque
    eq_table.rows[3].cells[0].text = 'Catégorie:'
    eq_table.rows[3].cells[1].text = intervention.demande.equipement.categorie.nom if intervention.demande.equipement.categorie else '-'
    eq_table.rows[4].cells[0].text = 'Bureau:'
    eq_table.rows[4].cells[1].text = intervention.demande.equipement.bureau.nom if intervention.demande.equipement.bureau else '-'
    
    # Détails
    document.add_heading('Détails de l\'Intervention', level=1)
    document.add_paragraph(intervention.details)
    
    # Pièces
    if intervention.pieces.exists():
        document.add_heading('Pièces de Rechange', level=1)
        pieces_table = document.add_table(rows=intervention.pieces.count()+2, cols=4)
        pieces_table.style = 'Light Grid Accent 1'
        
        pieces_table.rows[0].cells[0].text = 'Nom'
        pieces_table.rows[0].cells[1].text = 'Prix Unitaire'
        pieces_table.rows[0].cells[2].text = 'Quantité'
        pieces_table.rows[0].cells[3].text = 'Total'
        
        for i, piece in enumerate(intervention.pieces.all(), 1):
            pieces_table.rows[i].cells[0].text = piece.nom
            pieces_table.rows[i].cells[1].text = f"{piece.prix_unitaire} DA"
            pieces_table.rows[i].cells[2].text = str(piece.quantite)
            pieces_table.rows[i].cells[3].text = f"{piece.cout_total()} DA"
        
        last_row = pieces_table.rows[-1]
        last_row.cells[2].text = 'TOTAL:'
        last_row.cells[3].text = f"{intervention.cout_total_pieces()} DA"
    
    # --- VISAS (En bas du document) ---
    document.add_paragraph("\n") # Espacement avant les signatures
    visa_table = document.add_table(rows=1, cols=2)
    visa_table.autofit = True
    visa_table.style = None
    # Visa Informaticien (Gauche)
    v_info = visa_table.rows[0].cells[0].paragraphs[0]
    v_info.text = "VISA DE L'INFORMATICIEN"
    v_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    v_info.runs[0].bold = True

    # Visa Responsable (Droite)
    v_resp = visa_table.rows[0].cells[1].paragraphs[0]
    v_resp.text = "VISA DE RESPONSABLE DE LA CELLULE"
    v_resp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    v_resp.runs[0].bold = True

    # Sauvegarder
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="intervention_{intervention.pk}.docx"'
    return response


# ============= CONSULTATION LOGS (ADMIN) =============

@login_required
@user_passes_test(is_admin)
def admin_liste_logs(request):
    """Liste de tous les logs avec filtres"""
    logs = LogAction.objects.select_related('utilisateur').order_by('-date_action')
    
    # Appliquer les filtres
    form = FiltreLogForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('utilisateur'):
            logs = logs.filter(utilisateur=form.cleaned_data['utilisateur'])
        if form.cleaned_data.get('action'):
            logs = logs.filter(action=form.cleaned_data['action'])
        if form.cleaned_data.get('date_debut'):
            logs = logs.filter(date_action__gte=form.cleaned_data['date_debut'])
        if form.cleaned_data.get('date_fin'):
            from datetime import datetime, time
            date_fin = datetime.combine(form.cleaned_data['date_fin'], time.max)
            logs = logs.filter(date_action__lte=date_fin)
        if form.cleaned_data.get('recherche'):
            logs = logs.filter(details__icontains=form.cleaned_data['recherche'])
    
    # Pagination
    
    paginator = Paginator(logs, 20) 
    page_number = request.GET.get('page')
    logs_page = paginator.get_page(page_number)
    
    # Statistiques
    stats = {
        'total': logs.count(),
        'aujourd_hui': logs.filter(date_action__date=timezone.now().date()).count(),
        'cette_semaine': logs.filter(date_action__gte=timezone.now() - timezone.timedelta(days=7)).count(),
    }
    
    context = {
        'logs': logs_page,
        'form': form,
        'stats': stats,
    }
    return render(request, 'maintenance/admin/liste_logs.html', context)


@login_required
@user_passes_test(is_admin)
def admin_export_logs_csv(request):
    """Export des logs en CSV"""
    logs = LogAction.objects.select_related('utilisateur').order_by('-date_action')
    
    # Appliquer les mêmes filtres que la liste
    form = FiltreLogForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('utilisateur'):
            logs = logs.filter(utilisateur=form.cleaned_data['utilisateur'])
        if form.cleaned_data.get('action'):
            logs = logs.filter(action=form.cleaned_data['action'])
        if form.cleaned_data.get('date_debut'):
            logs = logs.filter(date_action__gte=form.cleaned_data['date_debut'])
        if form.cleaned_data.get('date_fin'):
            from datetime import datetime, time
            date_fin = datetime.combine(form.cleaned_data['date_fin'], time.max)
            logs = logs.filter(date_action__lte=date_fin)
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="logs_{datetime.now().strftime("%Y%m%d_%H%M")}.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Date', 'Heure', 'Utilisateur', 'Rôle', 'Action', 'Type Objet', 'ID Objet', 'Détails', 'IP'])
    
    for log in logs:
        writer.writerow([
            log.date_action.strftime('%Y-%m-%d'),
            log.date_action.strftime('%H:%M:%S'),
            log.utilisateur.get_full_name() if log.utilisateur else 'Système',
            log.utilisateur.get_role_display() if log.utilisateur else '-',
            log.get_action_display(),
            log.type_objet or '-',
            log.objet_id or '-',
            log.details[:200],
            log.adresse_ip or '-',
        ])
    
    return response


@login_required
@user_passes_test(is_admin)
def admin_export_logs_pdf(request):
    """Export des logs en PDF"""
    logs = LogAction.objects.select_related('utilisateur').order_by('-date_action')  # Limite 100 pour PDF
    
    # Appliquer les filtres
    form = FiltreLogForm(request.GET)
    if form.is_valid():
        if form.cleaned_data.get('utilisateur'):
            logs = logs.filter(utilisateur=form.cleaned_data['utilisateur'])[:100]
        if form.cleaned_data.get('action'):
            logs = logs.filter(action=form.cleaned_data['action'])[:100]
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
    elements = []
    
    styles = getSampleStyleSheet()
    title = Paragraph("RAPPORT DES LOGS D'ACTIONS", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 20))
    
    # Info
    info = Paragraph(f"Généré le {timezone.now().strftime('%d/%m/%Y à %H:%M')}", styles['Normal'])
    elements.append(info)
    elements.append(Spacer(1, 20))
    
    # Données
    data = [['Date', 'Utilisateur', 'Action', 'Détails']]
    
    for log in logs:
        data.append([
            log.date_action.strftime('%d/%m/%Y %H:%M'),
            log.utilisateur.get_full_name() if log.utilisateur else 'Système',
            log.get_action_display(),
            (log.details[:60] + '...') if len(log.details) > 60 else log.details,
        ])
    
    table = Table(data, colWidths=[4*cm, 5*cm, 6*cm, 10*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
    ]))
    
    elements.append(table)
    doc.build(elements)
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="logs_{datetime.now().strftime("%Y%m%d_%H%M")}.pdf"'
    return response